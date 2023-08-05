#an integration is a plugin that executes a binary on the operating system that this agent is running on and turns its output into Nodes that are ingested by Secquiry

import collablio.node as cnode
import collablio.client as cclient
import filereadutils
import logger

import json
import docx
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_BREAK

import random
import sys
import re
import os
import time

import copy
import datetime
import uuid
import traceback

import base64
import io

######################################################
'''
Notes:
* There needs to be a ReportTemplate.docx in the current folder, 
* custom styles should be saved into the docx, along with headers, footers, a title page
* the auto generated content will be appended to anything that is in the template
'''
######################################################

REPORT_TMP_DIR = '/tmp/'
START_HEADING_LEVEL = 0
MAX_HEADING_LEVEL = 9

# https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
def applyStylesToRun(style, run):
    if not style:
        return
    if 'bold' in style:
        run.bold = True
    if 'italic' in style:
        run.italic = True


def initialiseDocStyles(doc):
    if 'Code Block' not in doc.styles:
        codeBlockStyle = doc.styles.add_style('Code Block',docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        codeBlockStyle.base_style = doc.styles['Body Text']
        codeBlockStyle.font.name = 'Consolas'
    
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def applyStylesToParagraph(style, paragraph, doc, startHeadingLevel):
    if not style:
        return
    docStyles = doc.styles
    pformat = paragraph.paragraph_format
    if 'code-block' in style:
        testTable = doc.add_table(1,1)
        testPara = testTable.cell(0,0).paragraphs[0]
        testTable.style = docStyles['Medium Shading 1 Accent 1'] #need to change this to use style Name (not style ID) to get rid of warning
        testTable.style.font.bold = False
        #testTable.style.font.italic = True
        testPara.style.font.bold = False
        testPara.paragraph_format.space_before = Cm(0.1)
        testPara.paragraph_format.line_spacing = 1

        for run in paragraph.runs:
            testPara.add_run(run.text, run.style)

        testPara.style = docStyles['Code Block']
        delete_paragraph(paragraph)
        trailingPara = doc.add_paragraph()
        trailingPara.paragraph_format.space_after = Cm(0)

        return # we've nuked paragraph, nothing else can reference it at this point

    if 'list' in style:
        listItemType = style['list']
        paragraph.style = docStyles['List Bullet']
        pformat.left_indent = Cm(1)
        pformat.space_before = Cm(0)
        pformat.space_after = Cm(0)

    if 'header' in style:
        headingLevel = int(style['header'])
        headingLevel += startHeadingLevel
        if headingLevel > MAX_HEADING_LEVEL:
            headingLevel = MAX_HEADING_LEVEL
        paragraph.style = docStyles['Heading '+str(headingLevel)]

    if 'blockquote' in style:
        paragraph.style = docStyles['Quote']


    #ensure this is last
    if 'indent' in style:
        pformat.left_indent += Cm(int(style['indent']))
        #run.text = '\t'*int(style['indent'])+run.text

STYLES_WITH_SEPARATE_PARAGRAPHS = ['list']

def isStyleWithItsOwnParagraph(style):
    if not style:
        return False
    for key in list(style.keys()):
        if key in STYLES_WITH_SEPARATE_PARAGRAPHS:
            return True
    #print('key '+key+' is not in STYLES_WITH_SEPARATE_PARAGRAPHS')
    return False


# formatting attributes in Quill are generally applied to the last operation's inserted line
# some attributes like bold, italic, strikethrough are applied to the line that follows (forward-applied attributes):
FWD_APPLIED_ATTRIBS = ['bold','italic']

# otherAttachments are things like tables, images to might be referenced in this quillOpsData
# its type is:  { "reference" : "unique_ref_string", "nodeUid":"uid_of_node" }
def convertQuillOpsToDOCX(quilljsondata, doc, curdepth, otherAttachments = None):

    quillOpsData = json.loads(quilljsondata)

    Lines = [{'content':'','formatting':''}]

    count = 1
    for op in quillOpsData['ops']:
        #print('\n\nLine #{} *********'.format(count))
        count = count+1
        #print(op)
        fwdAttrs = None
        if 'attributes' in op: 
            attrs = op['attributes']
            #print('attrs='+str(attrs))
            opHasFwdAppliedAttrib = False
            #print(list(attrs.keys()))
            for attrName in list(attrs.keys()):
                #print('checking FWD_APPLIED_ATTRIBS for '+attrName)
                opHasFwdAppliedAttrib = (opHasFwdAppliedAttrib or (attrName.lower() in FWD_APPLIED_ATTRIBS))
            if not opHasFwdAppliedAttrib:
                #print('No op attributes are in FWD_APPLY_ATTRIBS'+str(attrs))
                Lines[-1]['formatting'] = attrs
                #print('Lines[-1]: '+str(Lines[-1]))
            else:
                #print('Apply forward attribute'+str(attrs))
                fwdAttrs =  attrs

        #handle images
        if isinstance(op['insert'], dict):
            insertData = op['insert']
            if 'image' in insertData:
                #print('-------found image: '+insertData['image'])
                Lines.append({ 'image': insertData['image'], 'content': '', 'formatting' : {} })
                
        elif isinstance(op['insert'], str):
            startFromStrPos = 0
            while True:
                indexOfNextNewline = op['insert'].find('\n',startFromStrPos)
                if indexOfNextNewline < 0:
                    #print('no more newlines found')
                    Lines.append({ 'content': op['insert'][startFromStrPos:], 'formatting' : fwdAttrs })
                    break
                #print('newline found at: '+str(indexOfNextNewline))
                Lines.append({ 'content': op['insert'][startFromStrPos:indexOfNextNewline+1], 'formatting' : fwdAttrs })
                startFromStrPos = indexOfNextNewline+1
                if startFromStrPos >= len(op['insert']):
                    break

    cleanedLines = [Lines[0]]
    for iL in range(0, len(Lines)):
        line = Lines[iL]
        #print(line)
        #merge lines that are just {line:'\n','formatting':None} with previous line if previous line isn't a sole newline 
        if 'content' in line and 'content' in Lines[iL-1]:
            lineBefore = Lines[iL-1]['content']
            if line['content'] == '\n' and not line['formatting'] and lineBefore != '\n':
                cleanedLines[-1]['content'] += '\n'
                #print('dropping empty line')
                continue
            
        cleanedLines.append(line)
            
    endOfNoteLine = { 'content': '', 'formatting': 'ENDOFNOTESTYLE' }
    cleanedLines.append(endOfNoteLine)

    #for line2 in cleanedLines:
    #    print(str(line2['formatting'])+' '+line2['content'])



    curLineIndex = 0
    prevLineStyle = {}
    runStr = ''
    prevStyleIsItsOwnParagraph = False

    while curLineIndex < len(cleanedLines):

        p = doc.add_paragraph('')
        finishedParagraph = False

        while not finishedParagraph:
            
            curLine = cleanedLines[curLineIndex]
            curLineContent = curLine['content']
            curLineStyle = curLine['formatting']

            prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)

            #print('*'*10+'\nprevStyle='+str(prevLineStyle)+' , curStyle='+str(curLineStyle)+', curLineContent=[['+curLineContent+']]')

            # the approach is to aggregate lines of the same style into the one paragraph
            # create a new paragraph once the style changes
            
            if ((curLineIndex >= len(cleanedLines) - 1) or ((curLineStyle != prevLineStyle) and runStr) or prevStyleIsItsOwnParagraph):
                #print('  - in add_run area')
                runEndsWithNewline = runStr.endswith('\n')
                if runEndsWithNewline:
                    runStr = runStr.strip('\n') #runStr[:-1] #get rid of trailing newline because space is automatically added between paragraphs
                run = p.add_run(runStr)
                                        
                if prevLineStyle:
                    applyStylesToRun(prevLineStyle, run)
                    
                #prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)
                #if prevStyleIsItsOwnParagraph:
                #    print('* '+str(prevLineStyle)+' is its own paragraph')
                finishedParagraph = (runEndsWithNewline or prevStyleIsItsOwnParagraph or (curLineIndex >= len(cleanedLines) - 1))
                if finishedParagraph:
                    applyStylesToParagraph(prevLineStyle, p, doc, curdepth)

                if 'image' in curLine:
                    try:
                        #{"insert":{"image":"data:image/png;base64,iVBORw0KGgoAAAANSUhEU...    
                        imgrun = doc.add_paragraph().add_run().add_picture(io.BytesIO(base64.b64decode(curLine['image'].split(',')[1])), width=Cm(17))
                    except:
                        print('an error occurred while decoding the base64 image')
                        traceback.print_exc()

                runStr = curLineContent #commence a new run string, assigning the current line
            else:
                runStr += curLineContent #append this line to the current run string

            prevLineStyle = curLineStyle
            curLineIndex += 1


CMS = 360000

def convertJexcelJsonToDOCX(jexcelJsonString, doc, curdepth, otherAttachments = None):
    #print(jexcelJsonString)
    '''
    t  = doc1.add_table(2,2)
    t.columns[0].width = int(1.5*INCHES)
    t.columns[1].width = int(5.25*INCHES)
    #t.cell(TR_RISK,TC_LABEL).text = LABEL_RISK
    #t.cell(TR_RISK,TC_VALUE).paragraphs[0].add_run().add_picture(RISK_IMAGES[findingData[FIELD_RISK]])
    
    '''
    try:
        jsonObj = json.loads(jexcelJsonString)
        #jsonObj { jsondata: '<serialisedJsonArray>', jsoncols: '<serialisedJsonArray>'}
        jsonRows = json.loads(jsonObj['jsondata'])
        jsonCols = json.loads(jsonObj['jsoncols']) if 'jsoncols' in jsonObj else []
        
        if len(jsonRows) < 1:
            return
        
        numOfColsInTable = len(jsonCols)
        numOfRowsInTable = len(jsonRows)+1 #need to add the header row
        
        tbl = doc.add_table(numOfRowsInTable, numOfColsInTable) #TODO: table style is 3rd argument
        
        colWidths = []
        for i,colData in enumerate(jsonCols):
            colWidths.append(int(colData['width']))
            curCell = tbl.cell(0,i)
            curCell.text = colData['title']
            #curCell.paragraphs[0].runs[0].style.bold = True

        totalTableWidth = 15.0*CMS
        totalColWidthUnits = float(sum(colWidths))
        for i,col in enumerate(tbl.columns):
            col.width = int((colWidths[i]/totalColWidthUnits)*totalTableWidth)
        
        for i,rowData in enumerate(jsonRows):
            for j in range(0,numOfColsInTable):
                curCell = tbl.cell(i+1,j)
                curCell.text = rowData[j]
        
    except Exception as e:
        print('an exception occurred while parsing table json: '+str(e))
        traceback.print_exc()

def addAPageBreak(doc):
    doc.add_paragraph('').add_run('').add_break(WD_BREAK.PAGE)

def sortNodes(uidDict):
    return getNodeForUID(uidDict[cnode.PROP_UID])[cnode.PROP_LABEL]
    
# don't call this on the reportRootNode, ignore it and only create its child sections
def docx_generate(docx, reportTreeNode, depth):
    #print('generate')
    if not reportTreeNode:
        #print('calling generate on a None reportTreeNode')
        return
    
    if reportTreeNode[cnode.PROP_TYPE] not in [cnode.TYPE_SECTION, cnode.TYPE_NOTE, cnode.TYPE_TEXT, cnode.TYPE_IMAGE, cnode.TYPE_TABLE]:
        return

    #assumes that the previous section has ended with a page break
    heading = 'Untitled Section'
    try:
        heading = reportTreeNode[cnode.PROP_LABEL]
    except:
        print('error when attempting to retrieve label for section '+reportTreeNode[cnode.PROP_UID])

    #putting "[N]" in the label is currently a dodgy hack to order the sections/notes, so remove the prefix from the label
    match = re.search('^\\[[a-z0-9]+\\][\\s]+', heading)
    if(match):
        heading = heading[len(match.group(0)):]
    #need to replace it with a proper x.y.z numbering scheme, by passing the parent section's number as a string to this function
    
    headingLevel = MAX_HEADING_LEVEL if (depth > MAX_HEADING_LEVEL) else depth

    p = docx.add_paragraph(heading, 'Heading '+str(headingLevel))

    #render table first??
    # or just require {{reference}} to tables and leave it up to user?
    if cnode.PROP_TEXTDATA in reportTreeNode:
        if reportTreeNode[cnode.PROP_TYPE] == cnode.TYPE_NOTE:
            convertQuillOpsToDOCX(reportTreeNode[cnode.PROP_TEXTDATA], docx, depth)
        elif reportTreeNode[cnode.PROP_TYPE] == cnode.TYPE_TABLE:
            convertJexcelJsonToDOCX(reportTreeNode[cnode.PROP_TEXTDATA], docx, depth)
        addAPageBreak(docx)

    if cnode.PROP_CHILDLIST in reportTreeNode:
        reportTreeNode[cnode.PROP_CHILDLIST].sort(key=sortNodes)
        for childSectionUID in reportTreeNode[cnode.PROP_CHILDLIST]:
            docx_generate(docx, getNodeForUID(childSectionUID[cnode.PROP_UID]), depth+1)


NODE_INDEX = {'_':''}

def getNodeForUID(uid):
    global NODE_INDEX    
    node = NODE_INDEX[uid] if uid in NODE_INDEX else None
    return node

def storeNode(node):
    global NODE_INDEX
    NODE_INDEX[node[cnode.PROP_UID]] = node

def clearNodeIndex():
    global NODE_INDEX
    NODE_INDEX = {'_':''}
    
'''
public class QueryNodesPostData
{
    public List<string> uids {get; set;}
    public string field {get; set;}
    public string op {get; set;}
    public string val {get; set;}
    public int depth {get; set;}
    public string type {get; set;}
}
'''

def generateReportForReportNode(reportRootNodeUID, client):
    try:
        jsonResponse =  client.fetchNodes(uid = reportRootNodeUID, body = True) #fetchNodes(querystring)

        for nodeResult in jsonResponse['nodes']:
            storeNode(nodeResult)
            #print(str(nodeResult))

        reportRootNode = getNodeForUID(reportRootNodeUID)
        
        newdoc = docx.Document()
        initialiseDocStyles(newdoc)
        
        #Initialise page margins for maximum text and image placement
        section = newdoc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(22.4)
        section.bottom_margin = Mm(22.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
        
        sectionDepth = 1

        for section in reportRootNode[cnode.PROP_CHILDLIST]:
            docx_generate(newdoc, getNodeForUID(section[cnode.PROP_UID]), sectionDepth)
        
        reportSaveName = REPORT_TMP_DIR+'report-'+str(int(time.time()))+str(uuid.uuid4()).replace('-','')[:4]+'.docx'
        newdoc.save(reportSaveName)
        #print(reportSaveName)
        return reportSaveName
            
    except Exception as e:
        logger.logEvent('an exception occurred while generating the report: '+str(e))
        logger.logEvent(traceback.format_exc())
        return ''


async def generate(metadata_dict):
    client = cclient.client #Client(metadata_dict['auth_token_hdr_val'])
    # fetch report root node and generate report docx file
    reportfile = generateReportForReportNode(metadata_dict['under_uid'], client)
    # upload report as file attachment node

    logger.logEvent(f'reportfile: {reportfile}')


    params = { 'parentid': metadata_dict['under_uid'] }

    # Create the form with simple fields
    form = cclient.MultiPartForm()
    form.add_field('type', 'file_upload')
    form.add_field('_p', json.dumps(params))

    # Add the file
    form.add_file('filedata', reportfile.split('/')[-1], fileHandle=open(reportfile, "rb"))

    client.createFileNode(form)
    '''
    # Build the request, including the byte-string
    # for the data to be posted.
    data = bytes(form)
    r = urllib.request.Request(COLLABLIO_HOST+'/upload', data=data) #  'http://127.0.0.1:9123'

    r.add_header('Content-type', form.get_content_type())
    r.add_header('Content-length', len(data))

    print()
    print('OUTGOING DATA:')
    for name, value in r.header_items():
        print('{}: {}'.format(name, value))

    #print(r.data.decode('utf-8'))

    #respStr = urllib.request.urlopen(r).read().decode('utf-8')
    respStr = executeHttpRequest(r).read().decode('utf-8')
    print('SERVER RESPONSE:')
    print(respStr)
    '''

    os.remove(reportfile)

    return {"status":"OK","detail":"Generated Successfully"}
    #return nl

