#!/usr/bin/env python3
#import http.server
#import socketserver
import json
import docx
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_BREAK

import random
import sys
import re
import os
import time

import copy
import datetime
import uuid
import urllib.request
import urllib.parse
import traceback

import io
import mimetypes
import base64

######################################################
'''
Notes:
* There needs to be a ReportTemplate.docx in the current folder, 
* custom styles should be saved into the docx, along with headers, footers, a title page
* the auto generated content will be appended to anything that is in the template

TODO:
make this a microservice/daemon that polls the database for report nodes,
    if a report node has an annotation (eg. 'generate') then do the auto generation, 
    upload the report as a file attachment to the report node
    delete the annotation (create a folder called 'completed jobs' and move it there or maybe just the recycle bin)
'''
######################################################

C_USAGE = '''Usage: {0} <projectname> <reportname>
        generate a docx report
'''

COLLABLIO_HOST = 'http://127.0.0.1:5000'

PROP_UID = "uid"
PROP_TYPE = "ty"
PROP_LABEL = "l"
PROP_DETAIL = "d"
PROP_TEXTDATA = "x"
PROP_CUSTOM = "c"
PROP_TIME = "t"
PROP_LASTMOD = "m"
PROP_BINARYDATA = "b"
PROP_EDITING = "e"
PROP_PARENTLIST = "in"
PROP_CHILDLIST = "out"
PROP_RELATIONS = "lnk"

TYPE_CLIENT = "Client"
TYPE_PROJECT = "Project"
TYPE_FOLDER = "Folder"
TYPE_HOST = "Host"
TYPE_PORT = "Port"
TYPE_TEXT = "Text"
TYPE_IMAGE = "Image"
TYPE_FILE = "File"
TYPE_NOTE = "Note"
TYPE_TABLE = "Table"
TYPE_ANNOTATION = "Annotation"
TYPE_TAG = "Tag"
TYPE_REPORT = "Report"
TYPE_SECTION = "Section"
TYPE_JOBREQ = "Job Request"


DAEMON_LISTENER_MODE = False

def parseArgs():
    global DAEMON_LISTENER_MODE
    data = {}
    try:
        num_args = len(sys.argv)
        data['projectname'] = sys.argv[1]
        data['reportname'] = sys.argv[2]
    except:
        print(C_USAGE.format(sys.argv[0]))

    #start in autogenerator listen mode
    DAEMON_LISTENER_MODE = not (('projectname' in data) and ('reportname' in data))
    return data

#del foo.bar
class Node:
    
    def __init__(self, _type):
        self.Type = _type
        self.Children = []
        self.Parents = []
        self.UID = ''
        self.Label = ''
        self.Detail = ''
        self.CustomData = ''
        self.TextData = ''
        
    def convert(self):
        apiFormatNode = {}
        apiFormatNode[PROP_UID] = self.UID
        apiFormatNode[PROP_TYPE] = self.Type
        apiFormatNode[PROP_LABEL] = self.Label
        apiFormatNode[PROP_DETAIL] = self.Detail
        apiFormatNode[PROP_TEXTDATA] = self.TextData
        #only need to specify the parent UID for the host nodes if they don't have a UID (i.e. a new insert)
        apiFormatNode[PROP_PARENTLIST] = [{PROP_UID : self.Parents[0]}] if (len(self.Parents) > 0) else []
        apiFormatNode[PROP_CHILDLIST] = [{PROP_UID : child.UID} for child in self.Children]
        return apiFormatNode
        

def executeHttpRequest(request):
    authreq = urllib.request.Request('http://127.0.0.1:5001/service/gettemptoken', data='null'.encode('utf8'), headers={'content-type': 'application/json'})
    response = urllib.request.urlopen(authreq)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    if 'token' not in jsonResponse:
        #raise Exception()
        jsonResponse = {'token':''}
    
    request.headers['Authorization'] = 'Bearer '+jsonResponse['token']
    return urllib.request.urlopen(request)

def fetchNodes(query):
    req = urllib.request.Request(COLLABLIO_HOST+"/nodes"+query)
    response = executeHttpRequest(req)
    #def fetchNodes(query):
    #response = urllib.request.urlopen(COLLABLIO_HOST+"/nodes"+query)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    if 'nodes' not in jsonResponse:
        raise Exception()
    return jsonResponse


def recursiveConvertNodesToAPIFormat(node, listToAddTheNodeTo):
    listToAddTheNodeTo.append(node.convert())
    if node.Children:
        for child in node.Children:
            recursiveConvertNodesToAPIFormat(child, listToAddTheNodeTo)
        




class MultiPartForm:
    """Accumulate the data to be used when posting a form."""

    def __init__(self):
        self.form_fields = []
        self.files = []
        # Use a large random byte string to separate
        # parts of the MIME data.
        self.boundary = uuid.uuid4().hex.encode('utf-8')
        return

    def get_content_type(self):
        return 'multipart/form-data; boundary={}'.format(
            self.boundary.decode('utf-8'))

    def add_field(self, name, value):
        """Add a simple field to the form data."""
        self.form_fields.append((name, value))

    def add_file(self, fieldname, filename, fileHandle,
                 mimetype=None):
        """Add a file to be uploaded."""
        body = fileHandle.read()
        if mimetype is None:
            mimetype = (
                mimetypes.guess_type(filename)[0] or
                'application/octet-stream'
            )
        self.files.append((fieldname, filename, mimetype, body))
        return

    @staticmethod
    def _form_data(name):
        return ('Content-Disposition: form-data; '
                'name="{}"\r\n').format(name).encode('utf-8')

    @staticmethod
    def _attached_file(name, filename):
        #return ('Content-Disposition: file; '
        return ('Content-Disposition: form-data; '
                'name="{}"; filename="{}"\r\n').format(
                    name, filename).encode('utf-8')

    @staticmethod
    def _content_type(ct):
        return 'Content-Type: {}\r\n'.format(ct).encode('utf-8')

    def __bytes__(self):
        """Return a byte-string representing the form data,
        including attached files.
        """
        buffer = io.BytesIO()
        boundary = b'--' + self.boundary + b'\r\n'

        # Add the form fields
        for name, value in self.form_fields:
            buffer.write(boundary)
            buffer.write(self._form_data(name))
            buffer.write(b'\r\n')
            buffer.write(value.encode('utf-8'))
            buffer.write(b'\r\n')

        # Add the files to upload
        for f_name, filename, f_content_type, body in self.files:
            buffer.write(boundary)
            buffer.write(self._attached_file(f_name, filename))
            buffer.write(self._content_type(f_content_type))
            buffer.write(b'\r\n')
            buffer.write(body)
            buffer.write(b'\r\n')

        buffer.write(b'--' + self.boundary + b'--\r\n')
        return buffer.getvalue()
        


START_HEADING_LEVEL = 0
MAX_HEADING_LEVEL = 9

# https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
def applyStylesToRun(style, run):
    if not style:
        return
    if 'bold' in style:
        run.bold = True
    if 'italic' in style:
        run.italic = True


def initialiseDocStyles(doc):
    if 'Code Block' not in doc.styles:
        codeBlockStyle = doc.styles.add_style('Code Block',docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        codeBlockStyle.base_style = doc.styles['Body Text']
        codeBlockStyle.font.name = 'Consolas'
    
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def applyStylesToParagraph(style, paragraph, doc, startHeadingLevel):
    if not style:
        return
    docStyles = doc.styles
    pformat = paragraph.paragraph_format
    if 'code-block' in style:
        testTable = doc.add_table(1,1)
        testPara = testTable.cell(0,0).paragraphs[0]
        testTable.style = docStyles['Medium Shading 1 Accent 1'] #need to change this to use style Name (not style ID) to get rid of warning
        testTable.style.font.bold = False
        #testTable.style.font.italic = True
        testPara.style.font.bold = False
        testPara.paragraph_format.space_before = Cm(0.1)
        testPara.paragraph_format.line_spacing = 1

        for run in paragraph.runs:
            testPara.add_run(run.text, run.style)

        testPara.style = docStyles['Code Block']
        delete_paragraph(paragraph)
        trailingPara = doc.add_paragraph()
        trailingPara.paragraph_format.space_after = Cm(0)

        return # we've nuked paragraph, nothing else can reference it at this point

    if 'list' in style:
        listItemType = style['list']
        paragraph.style = docStyles['List Bullet']
        pformat.left_indent = Cm(1)
        pformat.space_before = Cm(0)
        pformat.space_after = Cm(0)

    if 'header' in style:
        headingLevel = int(style['header'])
        headingLevel += startHeadingLevel
        if headingLevel > MAX_HEADING_LEVEL:
            headingLevel = MAX_HEADING_LEVEL
        paragraph.style = docStyles['Heading '+str(headingLevel)]

    if 'blockquote' in style:
        paragraph.style = docStyles['Quote']


    #ensure this is last
    if 'indent' in style:
        pformat.left_indent += Cm(int(style['indent']))
        #run.text = '\t'*int(style['indent'])+run.text

STYLES_WITH_SEPARATE_PARAGRAPHS = ['list']

def isStyleWithItsOwnParagraph(style):
    if not style:
        return False
    for key in list(style.keys()):
        if key in STYLES_WITH_SEPARATE_PARAGRAPHS:
            return True
    #print('key '+key+' is not in STYLES_WITH_SEPARATE_PARAGRAPHS')
    return False


# formatting attributes in Quill are generally applied to the last operation's inserted line
# some attributes like bold, italic, strikethrough are applied to the line that follows (forward-applied attributes):
FWD_APPLIED_ATTRIBS = ['bold','italic']

# otherAttachments are things like tables, images to might be referenced in this quillOpsData
# its type is:  { "reference" : "unique_ref_string", "nodeUid":"uid_of_node" }
def convertQuillOpsToDOCX(quilljsondata, doc, curdepth, otherAttachments = None):

    quillOpsData = json.loads(quilljsondata)

    Lines = [{'content':'','formatting':''}]

    count = 1
    for op in quillOpsData['ops']:
        #print('\n\nLine #{} *********'.format(count))
        count = count+1
        #print(op)
        fwdAttrs = None
        if 'attributes' in op: 
            attrs = op['attributes']
            #print('attrs='+str(attrs))
            opHasFwdAppliedAttrib = False
            #print(list(attrs.keys()))
            for attrName in list(attrs.keys()):
                #print('checking FWD_APPLIED_ATTRIBS for '+attrName)
                opHasFwdAppliedAttrib = (opHasFwdAppliedAttrib or (attrName.lower() in FWD_APPLIED_ATTRIBS))
            if not opHasFwdAppliedAttrib:
                #print('No op attributes are in FWD_APPLY_ATTRIBS'+str(attrs))
                Lines[-1]['formatting'] = attrs
                #print('Lines[-1]: '+str(Lines[-1]))
            else:
                #print('Apply forward attribute'+str(attrs))
                fwdAttrs =  attrs

        #handle images
        if isinstance(op['insert'], dict):
            insertData = op['insert']
            if 'image' in insertData:
                #print('-------found image: '+insertData['image'])
                Lines.append({ 'image': insertData['image'], 'content': '', 'formatting' : {} })
                
        elif isinstance(op['insert'], str):
            startFromStrPos = 0
            while True:
                indexOfNextNewline = op['insert'].find('\n',startFromStrPos)
                if indexOfNextNewline < 0:
                    #print('no more newlines found')
                    Lines.append({ 'content': op['insert'][startFromStrPos:], 'formatting' : fwdAttrs })
                    break
                #print('newline found at: '+str(indexOfNextNewline))
                Lines.append({ 'content': op['insert'][startFromStrPos:indexOfNextNewline+1], 'formatting' : fwdAttrs })
                startFromStrPos = indexOfNextNewline+1
                if startFromStrPos >= len(op['insert']):
                    break

    cleanedLines = [Lines[0]]
    for iL in range(0, len(Lines)):
        line = Lines[iL]
        #print(line)
        #merge lines that are just {line:'\n','formatting':None} with previous line if previous line isn't a sole newline 
        if 'content' in line and 'content' in Lines[iL-1]:
            lineBefore = Lines[iL-1]['content']
            if line['content'] == '\n' and not line['formatting'] and lineBefore != '\n':
                cleanedLines[-1]['content'] += '\n'
                #print('dropping empty line')
                continue
            
        cleanedLines.append(line)
            
    endOfNoteLine = { 'content': '', 'formatting': 'ENDOFNOTESTYLE' }
    cleanedLines.append(endOfNoteLine)

    for line2 in cleanedLines:
        print(str(line2['formatting'])+' '+line2['content'])



    curLineIndex = 0
    prevLineStyle = {}
    runStr = ''
    prevStyleIsItsOwnParagraph = False

    while curLineIndex < len(cleanedLines):

        p = doc.add_paragraph('')
        finishedParagraph = False

        while not finishedParagraph:
            
            curLine = cleanedLines[curLineIndex]
            curLineContent = curLine['content']
            curLineStyle = curLine['formatting']

            prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)

            #print('*'*10+'\nprevStyle='+str(prevLineStyle)+' , curStyle='+str(curLineStyle)+', curLineContent=[['+curLineContent+']]')

            # the approach is to aggregate lines of the same style into the one paragraph
            # create a new paragraph once the style changes
            
            if ((curLineIndex >= len(cleanedLines) - 1) or ((curLineStyle != prevLineStyle) and runStr) or prevStyleIsItsOwnParagraph):
                #print('  - in add_run area')
                runEndsWithNewline = runStr.endswith('\n')
                if runEndsWithNewline:
                    runStr = runStr.strip('\n') #runStr[:-1] #get rid of trailing newline because space is automatically added between paragraphs
                run = p.add_run(runStr)
                                        
                if prevLineStyle:
                    applyStylesToRun(prevLineStyle, run)
                    
                #prevStyleIsItsOwnParagraph = isStyleWithItsOwnParagraph(prevLineStyle)
                #if prevStyleIsItsOwnParagraph:
                #    print('* '+str(prevLineStyle)+' is its own paragraph')
                finishedParagraph = (runEndsWithNewline or prevStyleIsItsOwnParagraph or (curLineIndex >= len(cleanedLines) - 1))
                if finishedParagraph:
                    applyStylesToParagraph(prevLineStyle, p, doc, curdepth)

                if 'image' in curLine:
                    try:
                        #{"insert":{"image":"data:image/png;base64,iVBORw0KGgoAAAANSUhEU...    
                        imgrun = doc.add_paragraph().add_run().add_picture(io.BytesIO(base64.b64decode(curLine['image'].split(',')[1])), width=Cm(17))
                    except:
                        print('an error occurred while decoding the base64 image')

                runStr = curLineContent #commence a new run string, assigning the current line
            else:
                runStr += curLineContent #append this line to the current run string

            prevLineStyle = curLineStyle
            curLineIndex += 1


CMS = 360000

def convertJexcelJsonToDOCX(jexcelJsonString, doc, curdepth, otherAttachments = None):
    print(jexcelJsonString)
    '''
    t  = doc1.add_table(2,2)
    t.columns[0].width = int(1.5*INCHES)
    t.columns[1].width = int(5.25*INCHES)
    #t.cell(TR_RISK,TC_LABEL).text = LABEL_RISK
    #t.cell(TR_RISK,TC_VALUE).paragraphs[0].add_run().add_picture(RISK_IMAGES[findingData[FIELD_RISK]])
    
    '''
    try:
        jsonObj = json.loads(jexcelJsonString)
        #jsonObj { jsondata: '<serialisedJsonArray>', jsoncols: '<serialisedJsonArray>'}
        jsonRows = json.loads(jsonObj['jsondata'])
        jsonCols = json.loads(jsonObj['jsoncols']) if 'jsoncols' in jsonObj else []
        
        if len(jsonRows) < 1:
            return
        
        numOfColsInTable = len(jsonCols)
        numOfRowsInTable = len(jsonRows)+1 #need to add the header row
        
        tbl = doc.add_table(numOfRowsInTable, numOfColsInTable) #TODO: table style is 3rd argument
        
        colWidths = []
        for i,colData in enumerate(jsonCols):
            colWidths.append(int(colData['width']))
            curCell = tbl.cell(0,i)
            curCell.text = colData['title']
            #curCell.paragraphs[0].runs[0].style.bold = True

        totalTableWidth = 15.0*CMS
        totalColWidthUnits = float(sum(colWidths))
        for i,col in enumerate(tbl.columns):
            col.width = int((colWidths[i]/totalColWidthUnits)*totalTableWidth)
        
        for i,rowData in enumerate(jsonRows):
            for j in range(0,numOfColsInTable):
                curCell = tbl.cell(i+1,j)
                curCell.text = rowData[j]
        
    except Exception as e:
        print('an exception occurred while parsing table json: '+str(e))
        traceback.print_exc()

def addAPageBreak(doc):
    doc.add_paragraph('').add_run('').add_break(WD_BREAK.PAGE)

def sortNodes(uidDict):
    return getNodeForUID(uidDict[PROP_UID])[PROP_LABEL]
    
# don't call this on the reportRootNode, ignore it and only create its child sections
def generate(docx, reportTreeNode, depth):
    print('generate')
    if not reportTreeNode:
        print('calling generate on a None reportTreeNode')
        return
    
    if reportTreeNode[PROP_TYPE] not in [TYPE_SECTION, TYPE_NOTE, TYPE_TEXT, TYPE_IMAGE, TYPE_TABLE]:
        return

    #assumes that the previous section has ended with a page break
    heading = 'Untitled Section'
    try:
        heading = reportTreeNode[PROP_LABEL]
    except:
        print('error when attempting to retrieve label for section '+reportTreeNode[PROP_UID])

    #putting "[N]" in the label is currently a dodgy hack to order the sections/notes, so remove the prefix from the label
    match = re.search('^\\[[a-z0-9]+\\][\\s]+', heading)
    if(match):
        heading = heading[len(match.group(0)):]
    #need to replace it with a proper x.y.z numbering scheme, by passing the parent section's number as a string to this function
    
    headingLevel = MAX_HEADING_LEVEL if (depth > MAX_HEADING_LEVEL) else depth

    p = docx.add_paragraph(heading, 'Heading '+str(headingLevel))

    #render table first??
    # or just require {{reference}} to tables and leave it up to user?
    if PROP_TEXTDATA in reportTreeNode:
        if reportTreeNode[PROP_TYPE] == TYPE_NOTE:
            convertQuillOpsToDOCX(reportTreeNode[PROP_TEXTDATA], docx, depth)
        elif reportTreeNode[PROP_TYPE] == TYPE_TABLE:
            convertJexcelJsonToDOCX(reportTreeNode[PROP_TEXTDATA], docx, depth)
        addAPageBreak(docx)

    if PROP_CHILDLIST in reportTreeNode:
        reportTreeNode[PROP_CHILDLIST].sort(key=sortNodes)
        for childSectionUID in reportTreeNode[PROP_CHILDLIST]:
            generate(docx, getNodeForUID(childSectionUID[PROP_UID]), depth+1)


NODE_INDEX = {'_':''}

def getNodeForUID(uid):
    global NODE_INDEX    
    node = NODE_INDEX[uid] if uid in NODE_INDEX else None
    return node

def storeNode(node):
    global NODE_INDEX
    NODE_INDEX[node[PROP_UID]] = node

def clearNodeIndex():
    global NODE_INDEX
    NODE_INDEX = {'_':''}
    
'''
public class QueryNodesPostData
{
    public List<string> uids {get; set;}
    public string field {get; set;}
    public string op {get; set;}
    public string val {get; set;}
    public int depth {get; set;}
    public string type {get; set;}
}
'''

def fetchNodesPost(uids = [], field = PROP_LASTMOD, op = 'gt', val = '0', depth = 20, typ = ''):

    reqdata = { 'uids': uids, 'field': field, 'op': op, 'val': val, 'depth': depth, 'type': typ }
    print(json.dumps(reqdata))
    req = urllib.request.Request(url=COLLABLIO_HOST+'/nodes', data=bytes(json.dumps(reqdata), encoding='utf-8'))
    req.add_header('Content-Type', 'application/json')
    #response = urllib.request.urlopen(req)
    response = executeHttpRequest(req)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    if 'nodes' not in jsonResponse:
        raise Exception()
    return jsonResponse

def moveNodesPost(moveData):

    print(json.dumps(moveData))
    req = urllib.request.Request(url=COLLABLIO_HOST+'/move', data=bytes(json.dumps(moveData), encoding='utf-8'))
    req.add_header('Content-Type', 'application/json')
    #response = urllib.request.urlopen(req)
    response = executeHttpRequest(req)
    jsonResponse =  json.loads(response.read().decode('utf8'))
    return jsonResponse

def getRecycleBinFolderUID():
    querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}'.format(\
        '',\
        PROP_LABEL,\
        'eq',\
        urllib.parse.quote('Recycle Bin'),\
        20,\
        TYPE_FOLDER)
    print(COLLABLIO_HOST+"/nodes"+querystring)

    try:
        jsonResponse =  fetchNodes(querystring)
        return jsonResponse['nodes'][0][PROP_UID]
    except:
        print('recycle bin folder doesnt exist, creating a new one')

    try:
        newRecycleBinNode = Node(TYPE_FOLDER)
        newRecycleBinNode.Label = 'Recycle Bin'
        nodesToUpsert = []
        recursiveConvertNodesToAPIFormat(newRecycleBinNode, nodesToUpsert)
        serialisedJson = json.dumps(nodesToUpsert).encode('utf8')
        req = urllib.request.Request(COLLABLIO_HOST+'/upsert', data=serialisedJson, headers={'content-type': 'application/json'})
        #response = urllib.request.urlopen(req)
        response = executeHttpRequest(req)
        returnedUids = json.loads(response.read().decode('utf8'))
        return returnedUids[0]
    except Exception as e:
        print('an exception occurred while creating recycle bin: '+str(e))
        traceback.print_exc()


##################################################################################
## The main program
##################################################################################


argdata = parseArgs()

#QueryNodesGet(string uid = null, string field=null, string op=null, string val=null, int depth = 0, string type = null)


def generateReportForReportNode(reportRootNode):
    try:
        #query recursively to find all child nodes under report node
        querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}&body=true'.format(\
            reportRootNode[PROP_UID],\
            PROP_LASTMOD,\
            'gt',\
            0,\
            20,\
            '')
        print(COLLABLIO_HOST+"/nodes"+querystring)

        jsonResponse =  fetchNodes(querystring)

        for nodeResult in jsonResponse['nodes']:
            storeNode(nodeResult)
            print(str(nodeResult))


        newdoc = docx.Document()
        initialiseDocStyles(newdoc)
        
        #Initialise page margins for maximum text and image placement
        section = newdoc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(22.4)
        section.bottom_margin = Mm(22.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
        
        sectionDepth = 1

        for section in reportRootNode[PROP_CHILDLIST]:
            generate(newdoc, getNodeForUID(section[PROP_UID]), sectionDepth)
        
        reportSaveName = 'report-'+str(int(time.time()))+str(uuid.uuid4()).replace('-','')[:4]+'.docx'
        newdoc.save(reportSaveName)
        print(reportSaveName)
        return reportSaveName
            
    except Exception as e:
        print('an exception occurred while generating the report: '+str(e))
        traceback.print_exc()
        return ''




#query to find projectname data['projectname']

if not DAEMON_LISTENER_MODE:
    querystring = '?uid=&field={}&op={}&val={}&depth={}&type={}'.format(\
        PROP_LABEL,\
        'eq',\
        urllib.parse.quote(argdata['projectname']),\
        20,\
        TYPE_PROJECT)
    print(COLLABLIO_HOST+"/nodes"+querystring)

    try:
        jsonResponse =  fetchNodes(querystring)
        projectUID = ''
        for nodeResult in jsonResponse['nodes']:
            if nodeResult[PROP_LABEL] == argdata['projectname']:
                projectUID = nodeResult[PROP_UID]
                break
        print('located project '+projectUID)

        #query recursively to find the specified reportname under projectname
        querystring = '?uid={}&field={}&op={}&val={}&depth={}&type={}'.format(\
            projectUID,\
            PROP_LABEL,\
            'eq',\
            urllib.parse.quote(argdata['reportname']),\
            20,\
            TYPE_REPORT)
        print(COLLABLIO_HOST+"/nodes"+querystring)

        jsonResponse =  fetchNodes(querystring)

        reportRootNode = None
        for nodeResult in jsonResponse['nodes']:
            if nodeResult[PROP_LABEL] == argdata['reportname']:
                reportRootNode = nodeResult

        if not reportRootNode:
            print('unable to locate report: '+argdata['reportname'])
            exit(0)

        print('located report '+argdata['reportname'])

        generateReportForReportNode(reportRootNode)

    except Exception as e:
        print('an exception occurred while generating the report: '+str(e))
        traceback.print_exc()
    
    exit(0)




print('No project or filename arguments were provided. Running in Daemon Listener Mode')

#query to find all new reports since last check
# todo: save the last check time in persistent storage to avoid fetching every single report upon the process running

lastFetchTime = 0

uidsOfReports = set()

while True:
    querystring = '?uid=&field={}&op={}&val={}&depth={}&type={}'.format(\
        PROP_LASTMOD,\
        'gt',\
        '__LASTMODTIME__',\
        20,\
        TYPE_REPORT)
    print('*'*30)
    time.sleep(10)
    querystring = querystring.replace('__LASTMODTIME__',str(lastFetchTime))
    print(COLLABLIO_HOST+"/nodes"+querystring)    
    try:
        jsonResponse =  fetchNodes(querystring)
        print(json.dumps(jsonResponse))
        
        if not 'nodes' in jsonResponse:
            continue

        if 'timestamp' in jsonResponse:
            lastFetchTime = int(jsonResponse['timestamp'])
            print('(cur)lastfetchtime='+str(lastFetchTime)+', str(int(jsonResponse[timestamp])='+str(int(jsonResponse['timestamp'])))

        clearNodeIndex()
        
        #uidsOfReports = []
        for nodeResult in jsonResponse['nodes']:
            if nodeResult[PROP_TYPE] == TYPE_REPORT:
                #uidsOfReports.append(nodeResult[PROP_UID])
                uidsOfReports.add(nodeResult[PROP_UID])
                storeNode(nodeResult)

        # now query the database for report generation jobrequests pending for any of those report nodes 
        #jsonResponse = fetchNodesPost(uids = uidsOfReports, typ = TYPE_JOBREQ)
        jsonResponse = fetchNodesPost(uids = list(uidsOfReports), typ = TYPE_JOBREQ, depth = 1)
        
        for nodeResult in jsonResponse['nodes']:
            print(str(nodeResult))
            if (PROP_PARENTLIST in nodeResult) and (len(nodeResult[PROP_PARENTLIST]) > 0):
                jobReqParentUID = nodeResult[PROP_PARENTLIST][0][PROP_UID]
                reportNode = getNodeForUID(jobReqParentUID)
                if reportNode:
                    reportfile = generateReportForReportNode(reportNode)
                    if reportfile:
                        print('about to upload '+reportfile)
                        #apparently python3 urllib doesn't have builtin support for multipart/form-data
                        # there's an implementation here https://pymotw.com/3/urllib.request/

                        #data.append('filedata', input.files[0]);
                        #data.append('type', 'file_upload');
                        #data.append('_p', JSON.stringify(params));
                        #await fetch('upload', {
                        #	params.parentid = node[PROP_PARENTLIST][0].uid;
                        
                        params = { 'parentid': jobReqParentUID }

                        # Create the form with simple fields
                        form = MultiPartForm()
                        form.add_field('type', 'file_upload')
                        form.add_field('_p', json.dumps(params))

                        # Add the file
                        form.add_file('filedata', reportfile, fileHandle=open(reportfile, "rb"))

                        # Build the request, including the byte-string
                        # for the data to be posted.
                        data = bytes(form)
                        r = urllib.request.Request(COLLABLIO_HOST+'/upload', data=data) #  'http://127.0.0.1:9123'

                        r.add_header('Content-type', form.get_content_type())
                        r.add_header('Content-length', len(data))

                        print()
                        print('OUTGOING DATA:')
                        for name, value in r.header_items():
                            print('{}: {}'.format(name, value))

                        #print(r.data.decode('utf-8'))

                        #respStr = urllib.request.urlopen(r).read().decode('utf-8')
                        respStr = executeHttpRequest(r).read().decode('utf-8')
                        print('SERVER RESPONSE:')
                        print(respStr)

                        #move the jobrequest node to the recycle bin
                        rbFolderUID = getRecycleBinFolderUID()
                        print('Recycle Bin UID: '+rbFolderUID)
                        

                        moveData = { 'nodes': [nodeResult[PROP_UID]], 'parents': [jobReqParentUID], 'children': [], 'newparent': rbFolderUID }
                        responseForMove = moveNodesPost(moveData)
                        if responseForMove and 'error' in responseForMove:
                            print('error moving jobrequest to recyclebin '+str(responseForMove))
                        else:
                            #delete the local report file
                            os.remove(reportfile)
            
    except Exception as e:
        print('an exception occurred while fetching the report/job nodes: '+str(e))
        traceback.print_exc()


                
exit(0)


